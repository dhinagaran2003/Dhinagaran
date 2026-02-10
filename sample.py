from docx import Document

doc = Document()
doc.add_heading('3-Month Django & React Weekly Content Plan', level=1)

table = doc.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Day'
hdr_cells[1].text = 'Topic'
hdr_cells[2].text = 'Content (with Subtitles)'

weeks = [
    ("Monday", "Django", "Introduction\n• What is Django\n• Why use Django"),
    ("Wednesday", "React", "Introduction\n• What is React\n• Why React is popular"),
    ("Friday", "Django", "Setup\n• Install Python\n• Install Django"),

    ("Monday", "React", "Setup\n• Node.js & npm\n• Create React App"),
    ("Wednesday", "Django", "Project Basics\n• Create Project\n• Run Server"),
    ("Friday", "React", "JSX Basics\n• JSX Syntax\n• Components"),

    ("Monday", "Django", "Project Structure\n• settings.py\n• urls.py"),
    ("Wednesday", "React", "Components\n• Functional Components\n• Props"),
    ("Friday", "Django", "MVT Architecture\n• Model\n• View\n• Template"),

    ("Monday", "React", "State\n• useState Hook\n• Example"),
    ("Wednesday", "Django", "Models\n• Create Model\n• Database Table"),
    ("Friday", "React", "Events\n• onClick\n• Handling Events"),

    ("Monday", "Django", "Admin Panel\n• Superuser\n• Register Models"),
    ("Wednesday", "React", "Forms\n• Controlled Forms\n• Input Handling"),
    ("Friday", "Django", "Views\n• Function Based Views"),

    ("Monday", "React", "useEffect\n• Side Effects\n• API Call"),
    ("Wednesday", "Django", "Templates\n• HTML Integration\n• Template Tags"),
    ("Friday", "React", "Conditional Rendering\n• if/else\n• Ternary"),

    ("Monday", "Django", "URLs\n• Path\n• Include"),
    ("Wednesday", "React", "Lists & Keys\n• map()\n• Keys"),
    ("Friday", "Django", "Forms\n• Django Forms\n• Model Forms"),

    ("Monday", "React", "CSS in React\n• Inline CSS\n• External CSS"),
    ("Wednesday", "Django", "Authentication\n• Login\n• Logout"),
    ("Friday", "React", "Props vs State\n• Difference"),

    ("Monday", "Django", "ORM\n• CRUD Operations"),
    ("Wednesday", "React", "API Integration\n• Fetch\n• Axios"),
    ("Friday", "Django", "Static Files\n• CSS\n• JS"),

    ("Monday", "React", "Routing\n• React Router"),
    ("Wednesday", "Django", "Security\n• CSRF\n• Authentication"),
    ("Friday", "React", "Project Structure\n• Best Practices"),

    ("Monday", "Django", "Deployment\n• Basics\n• Server"),
    ("Wednesday", "React", "Build Process\n• npm build"),
    ("Friday", "Django", "Mini Project\n• CRUD App"),

    ("Monday", "React", "Mini Project\n• Simple App"),
    ("Wednesday", "Django", "Career Path\n• Backend\n• Full Stack"),
    ("Friday", "React", "Career Path\n• Frontend Developer"),
]

for day, topic, content in weeks:
    row_cells = table.add_row().cells
    row_cells[0].text = day
    row_cells[1].text = topic
    row_cells[2].text = content

file_path = "/mnt/data/3_Month_Django_React_Weekly_Plan.docx"
doc.save(file_path)

file_path
